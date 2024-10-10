import ell
from typing import List as TypedList, List
from pydantic import BaseModel, Field
from rich.table import Table
from rich.console import Console
from rich import box
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from itertools import zip_longest  # Add this import at the top of your file

class QuestionReview(BaseModel):
    outcome: str = Field(description="The outcome of the problem written as PASS or FAIL")

class Question(BaseModel):
    question: str = Field(description="The exam question.")
    answers: TypedList[str] = Field(description="The possible answers to the exam question.")
    answer_rationale: str = Field(description="The rationale for the correct answer.")
    right_answer: str = Field(description="The right answer.")

ell.init(verbose=False)

@ell.complex(model="gpt-4o-2024-08-06", response_format=Question, temperature=0.9)
def generate_exam_question(topic: str, excluding_existing_questions: List[str] = [], excluding_existing_answers: List[str] = []):
    """
An expert at crafting graduate-level exam questions possesses a deep understanding of the 
subject matter and pedagogical principles. They combine analytical thinking with creativity to 
design questions that assess both knowledge and critical thinking skills. This individual stays 
current with the latest research and industry trends, ensuring questions are relevant and challenging. 
They have a keen ability to create scenarios that mirror real-world complexities, forcing students to apply 
theoretical concepts practically. Their questions are clear, unambiguous, and carefully worded to avoid 
unintentional hints. They skillfully balance difficulty levels, incorporating a mix of straightforward and 
nuanced problems. Additionally, they excel at developing comprehensive rubrics for consistent evaluation, 
and can anticipate potential misinterpretations to refine question phrasing.

The answers should be a variety of different lengths, from 2 words to 100 words.
The answers could be quotations as the response from the medical professional.
Include the correct answer and a detailed rationale explaining why it is correct and why the other options are incorrect.

Do not start the question with "A 36-year-old" format. Mix it up.
"""
    return f"""
Generate a graduate level multiple choice question on {topic} that 
is not in the list of excluded questions: {excluding_existing_questions}. 
The question should not have the answer in the list of excluded answers: {excluding_existing_answers}.
Include the correct answer and a detailed rationale.
"""

@ell.complex(model="gpt-4o-2024-08-06", response_format=QuestionReview, temperature=0.1)
def check_question_quality(question : str):
    """
A person skilled at evaluating multiple-choice graduate-level exam answers possesses 
sharp analytical abilities and subject expertise. They excel at identifying clarity and 
precision in answer choices, ensuring each option directly relates to the question posed. 
This individual has a keen eye for subtle distinctions, quickly spotting vague or ambiguous 
language that could confuse test-takers. They're adept at recognizing when distractors are too 
similar or too dissimilar to the correct answer, maintaining an appropriate level of challenge. 
Their strong grasp of the subject matter allows them to verify the accuracy of each option and 
its relevance to the question. They're also attuned to potential misinterpretations, consistently 
refining answer choices to eliminate any unintended correct responses or misleading phrasing.
    """
    return f"Review the following question and answers: {question}."


def generate_exam_questions(topic: str, max_attempts: int = 10, count: int = 3) -> TypedList[Question]:
    acceptable_questions: TypedList[Question] = []
    number_of_attempts = 0
    number_of_questions_found = 0
    
    while len(acceptable_questions) < count and number_of_attempts < max_attempts:
        excluded_questions = [question.question for question in acceptable_questions]
        # create a list variable that includes a unique set of answers from the acceptable questions
        excluded_answers = [answer for question in acceptable_questions for answer in question.answers]
        question = generate_exam_question(topic, excluded_questions, excluded_answers)
        review = check_question_quality(question.parsed.question)
        
        if review.parsed.outcome == "PASS":
            acceptable_questions.append(question.parsed)
            number_of_questions_found += 1
            print(f"Found {number_of_questions_found} acceptable questions.")
        number_of_attempts += 1

    if len(acceptable_questions) < count:
        print(f"Unable to generate {count} acceptable questions. Returning {len(acceptable_questions)} questions.")

    return acceptable_questions

def pretty_print_questions(questions: TypedList[Question]):
    table = Table(title="Exam Questions", box=box.SQUARE, show_lines=True)
    table.add_column("Question", style="cyan", no_wrap=False)
    table.add_column("Answers", style="magenta", no_wrap=False)

    for question in questions:
        # convert question.answers to a string
        answers = "\n".join([f"{chr(9311 + i)}. {answer}" for i, answer in enumerate(question.answers)])
        table.add_row(question.question, answers)

    console = Console()
    console.print(table)

def generate_exam_word(questions_depression: TypedList[Question], questions_bipolar: TypedList[Question], filename: str):
    doc = Document()
    
    # Add title
    doc.add_heading('Psychiatry Board Exam Questions', 0)

    # Create custom styles
    styles = doc.styles
    correct_answer_style = styles.add_style('CorrectAnswer', WD_STYLE_TYPE.PARAGRAPH)
    correct_answer_style.font.bold = True
    correct_answer_style.font.color.rgb = RGBColor(0, 128, 0)  # Dark green

    rationale_style = styles.add_style('Rationale', WD_STYLE_TYPE.PARAGRAPH)
    rationale_style.font.italic = True

    # Function to add a single question to the document
    def add_question_to_doc(question, question_number):
        if question is None:
            return

        # Add question
        doc.add_paragraph(f"{question_number}. {question.question}")
        
        # Add answers
        for j, answer in enumerate(question.answers):
            doc.add_paragraph(f"   {chr(97 + j)}) {answer}")
        
        # Add correct answer
        doc.add_paragraph(f"Correct Answer: {question.right_answer}", style='CorrectAnswer')
        
        # Add rationale
        doc.add_paragraph("Rationale:", style='Rationale')
        doc.add_paragraph(question.answer_rationale)
        
        # Add space between questions
        doc.add_paragraph()

    # Interleave questions from both lists
    for i, (depression_q, bipolar_q) in enumerate(zip_longest(questions_depression, questions_bipolar), 1):
        if depression_q:
            doc.add_paragraph("Topic: Major Depressive Disorders and Associated Disorders", style='Heading 2')
            add_question_to_doc(depression_q, i * 2 - 1)
        
        if bipolar_q:
            doc.add_paragraph("Topic: Bipolar Disorders and Related Disorders", style='Heading 2')
            add_question_to_doc(bipolar_q, i * 2)

    # Save the document
    doc.save(filename)

def generate_and_save_exam(topic: str, max_attempts: int, count: int) -> TypedList[Question]:
    questions = generate_exam_questions(topic, max_attempts, count)
    return questions


prompt = """
Write a question with multiple choice answers and rationale for psychiatry students. 
The questions should be equivalent to board exam questions. 
The questions should have scenarios and should elicit critical thinking. 
The questions should cover assessment, diagnosis, differential diagnosis, treatment, 
side effects, epidemiology, prognosis, therapeutic communication, education, and 
dsm5 criteria for major depressive disorder and related disorders and 
bipolar disorders and related disorders. 
"""

differentiator1 = f"{prompt} The question should be about major depressive disorders and associated disorders."
differentiator2 = f"{prompt} The question should be about bipolar disorders and related disorders."

number_of_questions = 25
max_attempts = 50

# Update the main execution part
depression_questions = generate_and_save_exam(
    f"{prompt} {differentiator1}", 
    max_attempts, 
    number_of_questions
)
pretty_print_questions(depression_questions)

bipolar_questions = generate_and_save_exam(
    f"{prompt} {differentiator2}", 
    max_attempts, 
    number_of_questions
)
pretty_print_questions(bipolar_questions)

# Generate single Word document with both sets of questions
generate_exam_word(depression_questions, bipolar_questions, "psychiatry_board_exam.docx")
print("Exam saved to psychiatry_board_exam.docx")

